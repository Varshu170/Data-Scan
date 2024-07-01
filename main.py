import pytesseract
from PIL import Image
import PyPDF2
from PyPDF2 import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import docx


def ocr_image(image):
    return pytesseract.image_to_string(image, lang='eng')


def extract_images_from_pdf(pdf_path):
    images = []
    with open(pdf_path, "rb") as pdf_file:
        reader = PdfReader(pdf_file, strict=False)
        pages = len(reader.pages)
        print("Total pages:", pages)

        for page_num in range(pages):
            page = reader.pages[page_num]
            if '/XObject' in page['/Resources']:
                xobjects = page['/Resources']['/XObject'].get_object()
                for obj in xobjects.values():
                    obj = obj.get_object()
                    if obj['/Subtype'] == '/Image':
                        images.append(obj)

    print("Total images:", len(images))
    return images


def image_preprocessing(image):
    if isinstance(image, bytes):
        preprocessed_image = Image.open(io.BytesIO(image))
    elif isinstance(image, str):
        preprocessed_image = Image.open(image)
    elif isinstance(image, PyPDF2.generic.EncodedStreamObject):
        image_data = image.get_data()
        preprocessed_image = Image.open(io.BytesIO(image_data))
    else:
        preprocessed_image = image

    if isinstance(preprocessed_image, Image.Image):
        preprocessed_image = preprocessed_image.convert("RGB")  # Convert PIL Image to RGB mode

    return preprocessed_image


def create_searchable_file(input_path, output_path, output_format):
    if input_path.lower().endswith('.pdf'):
        images = extract_images_from_pdf(input_path)
    elif input_path.lower().endswith(('.png', '.jpg', '.jpeg')):
        images = [Image.open(input_path)]
    else:
        print("Unsupported file format. Please provide a PDF or image file.")
        return

    if output_format == "pdf":
        pdf = canvas.Canvas(output_path, pagesize=letter)
        pdf.setFont("Helvetica", 12)
        pdf.setFillColor(colors.black)

        for image in images:
            preprocessed_image = image_preprocessing(image)
            extracted_text = ocr_image(preprocessed_image)

            extracted_text = extracted_text.encode('utf-8', 'ignore').decode('utf-8')
            lines = extracted_text.split('\n')

            page_height = letter[1]
            line_height = 15
            current_y = page_height - line_height

            for line in lines:
                pdf.drawString(100, current_y, line)
                current_y -= line_height

                if current_y < line_height:
                    pdf.showPage()  # Create a new page
                    current_y = page_height - line_height

            pdf.showPage()  # Create a new page after drawing content of the current page

        pdf.save()
        print("Text extraction and conversion to PDF completed.")
    elif output_format == "word":
        doc = docx.Document()
        total_pages = 0  # Variable to store the total number of pages in the scanned PDF

        for image in images:
            preprocessed_image = image_preprocessing(image)
            extracted_text = ocr_image(preprocessed_image)

            extracted_text = extracted_text.encode('utf-8', 'ignore').decode('utf-8')
            lines = extracted_text.split('\n')

            for line in lines:
                paragraph = doc.add_paragraph(line)
                paragraph.space_before = docx.shared.Pt(0)  # Remove space before paragraph
                paragraph.space_after = docx.shared.Pt(0)  # Remove space after paragraph

            total_pages += 1

        doc.save(output_path)

        # Adjust the number of pages in the Word document
        doc = docx.Document(output_path)
        while len(doc.sections) < total_pages:
            doc.add_section()
        while len(doc.sections) > total_pages:
            doc.sections.pop(-1)
        doc.save(output_path)

        print("Text extraction and conversion to Word completed.")
    elif output_format == "text":
        with open(output_path, "w", encoding="utf-8") as text_file:
            for image in images:
                preprocessed_image = image_preprocessing(image)
                extracted_text = ocr_image(preprocessed_image)

                extracted_text = extracted_text.encode('utf-8', 'ignore').decode('utf-8')
                lines = extracted_text.split('\n')

                for line in lines:
                    text_file.write(line + "\n")

        print("Text extraction and conversion to text completed.")
    else:
        print("Invalid output format. Please choose 'pdf', 'word', or 'text'.")


input_path = r"C:\Users\tamiz\Downloads\Scanneddocs\Scanneddocs\scansmpl.pdf"

output_format = input("Enter the desired output format (pdf, word, text): ")
output_path = input("Enter the output file path: ")

create_searchable_file(input_path, output_path, output_format)
